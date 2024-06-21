import sys
import pyperclip
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

def parse_markdown(md_text):
    lines = md_text.split('\n')
    parsed_lines = []

    for line in lines:
        if line.startswith('# '):
            parsed_lines.append(('heading1', line[2:]))
        elif line.startswith('## '):
            parsed_lines.append(('heading2', line[3:]))
        elif line.startswith('### '):
            parsed_lines.append(('heading3', line[4:]))
        elif line.startswith('#### '):
            parsed_lines.append(('heading4', line[5:]))
        elif line.startswith('##### '):
            parsed_lines.append(('heading5', line[6:]))
        elif line.startswith('###### '):
            parsed_lines.append(('heading6', line[7:]))
        elif line.startswith('* '):
            parsed_lines.append(('list_item', line[2:]))
        elif re.match(r'\d+\.\s', line):
            parsed_lines.append(('numbered_item', re.sub(r'\d+\.\s', '', line)))
        else:
            parsed_lines.append(('paragraph', line))
    
    return parsed_lines

def apply_formatting(paragraph, text):
    # Split the text by bold, italic, and underline markers
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|_.*?_)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            bold_run = paragraph.add_run(part[2:-2])
            bold_run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            italic_run = paragraph.add_run(part[1:-1])
            italic_run.italic = True
        elif part.startswith('_') and part.endswith('_'):
            underline_run = paragraph.add_run(part[1:-1])
            underline_run.underline = True
        else:
            paragraph.add_run(part)

def add_to_document(doc, parsed_lines):
    for line_type, content in parsed_lines:
        if line_type == 'heading1':
            p = doc.add_heading(level=1)
            apply_formatting(p, content)
        elif line_type == 'heading2':
            p = doc.add_heading(level=2)
            apply_formatting(p, content)
        elif line_type == 'heading3':
            p = doc.add_heading(level=3)
            apply_formatting(p, content)
        elif line_type == 'heading4':
            p = doc.add_heading(level=4)
            apply_formatting(p, content)
        elif line_type == 'heading5':
            p = doc.add_heading(level=5)
            apply_formatting(p, content)
        elif line_type == 'heading6':
            p = doc.add_heading(level=6)
            apply_formatting(p, content)
        elif line_type == 'list_item':
            p = doc.add_paragraph(style='List Bullet')
            apply_formatting(p, content)
        elif line_type == 'numbered_item':
            p = doc.add_paragraph(style='List Number')
            apply_formatting(p, content)
        else:
            p = doc.add_paragraph()
            apply_formatting(p, content)

def main():
    if len(sys.argv) != 2:
        print("Usage: mdtodoc <output_doc_name>")
        sys.exit(1)

    output_doc_name = sys.argv[1] + '.docx'

    # Get the current string from the clipboard
    md_text = pyperclip.paste()

    # Parse the markdown text
    parsed_lines = parse_markdown(md_text)

    # Create a new Word document
    doc = Document()

    # Add parsed content to the document
    add_to_document(doc, parsed_lines)

    # Save the document
    doc.save(output_doc_name)
    print(f"Document saved as {output_doc_name}")

if __name__ == "__main__":
    main()